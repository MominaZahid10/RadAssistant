"""
RadAssist AI — Document Ingestion Service

THIS IS THE HEART OF PHASE 2.

When a doctor uploads a file or we fetch a medical article, this service
processes it through a 4-stage pipeline:

    ┌─────────┐    ┌─────────┐    ┌─────────┐    ┌─────────┐
    │  PARSE  │ →  │  CHUNK  │ →  │  EMBED  │ →  │  STORE  │
    │         │    │         │    │         │    │         │
    │ Extract │    │ Split   │    │ Convert │    │ Save to │
    │ text    │    │ into    │    │ to      │    │ Qdrant  │
    │ from    │    │ small   │    │ vectors │    │ + update│
    │ PDF/    │    │ pieces  │    │ (384    │    │ Postgres│
    │ DOCX/   │    │ with    │    │ numbers │    │ status  │
    │ IMG/    │    │ overlap │    │ each)   │    │         │
    │ TXT     │    │         │    │         │    │         │
    └─────────┘    └─────────┘    └─────────┘    └─────────┘

STAGE 1 — PARSE:
  Extracts raw text from any supported file format:
  - PDF: PyMuPDF (fitz) — C-based, very fast
  - DOCX: python-docx — reads Word documents
  - TXT/MD: Plain file read
  - Images: Pillow + pytesseract (OCR) — reads text from photos/scans

STAGE 2 — CHUNK:
  Splits the extracted text into small, overlapping pieces.
  WHY? Embedding models have a token limit (~256 tokens for MiniLM).
  Also, smaller chunks give more precise search results.

  Example (chunk_size=512, overlap=50):
  ┌────────────────────── Chunk 0 ──────────────────────┐
  │ "Pneumonia is an infection that inflames the air     │
  │  sacs in one or both lungs. The air sacs may fill   │
  │  with fluid or pus (purulent material), causing..." │
  └───────────────────────────────┬──────────────────────┘
                                  │ overlap (50 chars)
                    ┌─────────────┴────────────────────────┐
                    │ "...causing cough with phlegm or pus, │
                    │  fever, chills, and difficulty         │
                    │  breathing. On chest X-ray, pneumonia  │
                    │  appears as..."                        │
                    └──────────────────────────────────────┘
                                Chunk 1

STAGE 3 — EMBED:
  Converts each chunk into a 384-dimensional vector using the
  embedding model. Chunks about similar topics get similar vectors.

STAGE 4 — STORE:
  Saves vectors + text to Qdrant, updates PostgreSQL document status.
"""

import io
import re
import traceback

import fitz  # PyMuPDF — imported as 'fitz' (historical name)
from docx import Document as DocxDocument
from fastapi.concurrency import run_in_threadpool
from PIL import Image

from app.config import get_settings
from app.services.embedding import embedding_service
from app.services.qdrant_service import qdrant_service
from app.services.lexical_service import lexical_index

settings = get_settings()

# Below this many characters, a "successful" parse isn't worth embedding.
# Set high enough that error placeholders and stray page numbers can't
# masquerade as real content.
MIN_USABLE_TEXT_CHARS = 50


class ParseError(Exception):
    """
    Raised when a file cannot be turned into usable text.

    WHY A CUSTOM EXCEPTION?
    So the pipeline can distinguish "this file is bad" (a user problem we
    should report clearly) from "our code crashed" (a bug we should see).
    Returning an error *string* instead of raising is how bad data ends up
    embedded in the vector store — see parse_image() for the war story.
    """
    pass


# ══════════════════════════════════════════════════════════════
# STAGE 1 — DOCUMENT PARSING
# ══════════════════════════════════════════════════════════════
# Each parser extracts raw text from a specific file format.
# They all return a plain string — the rest of the pipeline
# doesn't care what format the original file was.
#
# CONTRACT: a parser either returns usable text, or raises ParseError.
# It never returns a placeholder / error message as if it were content.
# ══════════════════════════════════════════════════════════════


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    WHY PyMuPDF?
    - Built on the MuPDF C engine → 10x faster than pure Python alternatives
    - Handles digital PDFs perfectly (most medical guidelines/textbooks)
    - Preserves reading order even in multi-column layouts
    
    HOW IT WORKS:
    1. Opens the PDF from bytes (no need to save to disk)
    2. Iterates through every page
    3. Extracts text from each page, preserving paragraph structure
    4. Joins all pages with newlines
    
    Args:
        file_bytes: Raw bytes of the PDF file
        
    Returns:
        Extracted text as a single string
    """
    text_parts = []

    try:
        # fitz.open() can read from bytes via a stream — no temp file needed
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                # get_text("text") extracts plain text in reading order
                # "blocks" mode would give us more structure, but plain
                # text is sufficient for chunking and embedding
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
    except Exception as e:
        raise ParseError(f"Could not read this PDF: {e}") from e

    if not text_parts:
        raise ParseError(
            "This PDF contains no extractable text. It is most likely a "
            "scanned document — every page is an image. Convert it to PNG/TIFF "
            "and upload it as an image so it goes through OCR instead."
        )

    return "\n\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX (Word) file.
    
    HOW IT WORKS:
    1. Opens the DOCX from bytes (it's actually a ZIP file internally)
    2. Iterates through all paragraphs
    3. Joins paragraphs with newlines, skipping empty ones
    
    DOCX STRUCTURE:
    A Word document is a ZIP file containing XML. python-docx
    handles all the XML parsing for us. We just iterate paragraphs.
    
    Args:
        file_bytes: Raw bytes of the DOCX file
        
    Returns:
        Extracted text as a single string
    """
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ParseError(f"Could not read this DOCX file: {e}") from e

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Tables are common in clinical documents (protocols, dose charts)
    # and live outside doc.paragraphs — they'd be silently dropped otherwise.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        raise ParseError("This DOCX file contains no readable text.")

    return "\n\n".join(paragraphs)


def parse_text(file_bytes: bytes) -> str:
    """
    Extract text from a plain text or Markdown file.
    
    The simplest parser — just decode bytes to string.
    Handles both UTF-8 and Latin-1 encoding (covers 99% of text files).
    
    Args:
        file_bytes: Raw bytes of the text file
        
    Returns:
        The decoded text string
    """
    # Try UTF-8 first (most common), fall back to Latin-1 (never fails)
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def parse_image(file_bytes: bytes) -> str:
    """
    Extract text from an image using OCR (Optical Character Recognition).
    
    WHEN IS THIS USED?
    - Scanned medical documents (old paper records digitized as images)
    - Photos of handwritten notes
    - Screenshots of reports from other systems
    - Reference images with text annotations
    
    HOW IT WORKS:
    1. Opens the image with Pillow (PIL)
    2. Passes it to Tesseract OCR engine
    3. Tesseract recognizes text patterns and returns a string
    
    REQUIREMENTS:
    - Tesseract must be installed in the Docker container (see Dockerfile)
    - For best results, images should be clear and well-lit
    - Handwriting recognition is basic — works best with printed text

    ⚠️  WHY THIS RAISES INSTEAD OF RETURNING AN ERROR STRING:
    An earlier version returned "[OCR failed: ...]" on error. That string is
    longer than the minimum-length check downstream, so it sailed through the
    pipeline, got embedded into a vector, and was stored in Qdrant while the
    document was marked "completed". The knowledge base silently filled with
    garbage that would later surface as RAG search results.

    A failed parse must FAIL. ingest_document() catches the exception and
    records status="failed" with the real reason, which is what the status
    column exists for.

    Args:
        file_bytes: Raw bytes of the image file (PNG, JPG, TIFF, etc.)

    Returns:
        Extracted text from the image

    Raises:
        ParseError: If Tesseract is missing, the image is unreadable, or
                    OCR produces no usable text.
    """
    try:
        import pytesseract
    except ImportError as e:
        raise ParseError(
            "pytesseract is not installed. Add it to requirements.txt."
        ) from e

    try:
        image = Image.open(io.BytesIO(file_bytes))

        # Convert to RGB if necessary (some formats like PNG have alpha channel)
        if image.mode != "RGB":
            image = image.convert("RGB")

        # Run OCR — pytesseract shells out to the Tesseract binary
        text = pytesseract.image_to_string(image)

    except pytesseract.TesseractNotFoundError as e:
        raise ParseError(
            "The Tesseract OCR engine is not installed in this container. "
            "Install it with: apt-get install -y tesseract-ocr tesseract-ocr-eng"
        ) from e
    except Exception as e:
        raise ParseError(f"OCR failed on this image: {e}") from e

    text = text.strip()
    if not text:
        raise ParseError(
            "OCR ran successfully but found no readable text in this image. "
            "The image may be blank, too low-resolution, or contain only "
            "graphics with no text."
        )

    return text


def parse_file(file_bytes: bytes, file_type: str) -> str:
    """
    Route a file to the correct parser based on its type.
    
    This is the ENTRY POINT for Stage 1. The rest of the pipeline
    calls this function and gets back plain text regardless of format.
    
    Args:
        file_bytes: Raw bytes of the uploaded file
        file_type: File extension without dot (e.g., "pdf", "docx", "png")
        
    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If the file type is not supported
    """
    parsers = {
        "pdf": parse_pdf,
        "docx": parse_docx,
        "txt": parse_text,
        "md": parse_text,
        # Image formats — all go through OCR
        "png": parse_image,
        "jpg": parse_image,
        "jpeg": parse_image,
        "tiff": parse_image,
        "bmp": parse_image,
    }
    
    parser = parsers.get(file_type.lower())
    if not parser:
        raise ParseError(
            f"Unsupported file type: '{file_type}'. "
            f"Supported: {', '.join(parsers.keys())}"
        )

    return parser(file_bytes)


# ══════════════════════════════════════════════════════════════
# STAGE 2 — TEXT CHUNKING
# ══════════════════════════════════════════════════════════════
# Splits extracted text into smaller pieces optimized for embedding
# and retrieval. This is implemented from scratch (no LangChain
# dependency) using a recursive splitting strategy.
# ══════════════════════════════════════════════════════════════


def chunk_text(
    text: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[str]:
    """
    Split text into overlapping chunks using recursive character splitting.
    
    STRATEGY (Recursive Splitting):
    We try to split at the BEST boundary first, falling back to worse ones:
    1. Double newline (paragraph break) — BEST, preserves complete paragraphs
    2. Single newline (line break) — good, preserves sentences
    3. Period + space (sentence end) — acceptable, keeps sentence integrity
    4. Space (word boundary) — fallback, at least doesn't break words
    5. Any character — last resort, only for very long words (rare)
    
    WHY THIS ORDER?
    Medical text has meaning at the paragraph level. A paragraph about
    "pneumonia findings" should ideally stay in ONE chunk. If we split
    at character boundaries, we might cut "pneu|monia" in half.
    
    Args:
        text: The full extracted text to split
        chunk_size: Max characters per chunk (default from config: 512)
        chunk_overlap: Overlap between chunks (default from config: 50)
        
    Returns:
        List of text chunks, each at most chunk_size characters
    """
    # ⚠️  `is None`, NOT `or` — 0 is a falsy but MEANINGFUL value here.
    # `chunk_overlap or settings.CHUNK_OVERLAP` silently rewrites an explicit
    # chunk_overlap=0 (a legitimate "no overlap" request) into the configured
    # default, so overlap could never be disabled. The same bug let
    # chunk_size=0 fall through to 512 and skip the validation below.
    chunk_size = settings.CHUNK_SIZE if chunk_size is None else chunk_size
    chunk_overlap = settings.CHUNK_OVERLAP if chunk_overlap is None else chunk_overlap

    # ── Guard against a config that can't terminate ──────────
    # The character-level fallback below steps by (chunk_size - chunk_overlap).
    # If overlap >= size, that step is 0 or negative → range() either yields
    # nothing or loops forever. Fail loudly at the call site instead.
    if chunk_size <= 0:
        raise ValueError(f"CHUNK_SIZE must be positive, got {chunk_size}")
    if chunk_overlap < 0:
        raise ValueError(f"CHUNK_OVERLAP cannot be negative, got {chunk_overlap}")
    if chunk_overlap >= chunk_size:
        raise ValueError(
            f"CHUNK_OVERLAP ({chunk_overlap}) must be smaller than "
            f"CHUNK_SIZE ({chunk_size}), otherwise chunking cannot advance."
        )

    # Clean the text — normalize whitespace, remove excessive blank lines
    text = text.strip()
    if not text:
        return []
    
    # If the entire text fits in one chunk, just return it
    if len(text) <= chunk_size:
        return [text]
    
    # Separators in order of preference (best → worst)
    separators = ["\n\n", "\n", ". ", " ", ""]
    
    chunks = []
    _recursive_split(text, separators, chunk_size, chunk_overlap, chunks)
    
    # Filter out empty/whitespace-only chunks
    chunks = [c.strip() for c in chunks if c.strip()]
    
    return chunks


def _tail_overlap(text: str, chunk_overlap: int) -> str:
    """
    Take the last `chunk_overlap` characters of `text`, snapped forward to the
    nearest word boundary.

    WHY THE SNAP MATTERS:
    A naive `text[-chunk_overlap:]` slices blindly by character count, so a
    chunk can begin mid-word. Real example from the seeded knowledge base,
    where "...must not be missed on chest imaging" became:

        "e missed on chest imaging.\\n\\nTypes of Pneumothorax:..."

    Two concrete costs:
      1. The orphaned "e" is tokenised as a meaningless fragment, adding noise
         to a 384-dimensional vector that should represent medical meaning.
      2. These chunks are shown verbatim to the radiologist as the evidence
         behind a generated finding. A citation that opens mid-word reads as
         broken software, which is corrosive for a clinical decision-support
         tool whose entire value proposition is traceable, trustworthy sources.

    If the window contains no whitespace at all (one very long token, e.g. a
    chemical name or an OCR artefact), the raw slice is returned — trimming to
    nothing would lose the overlap entirely.
    """
    window = text[-chunk_overlap:]

    match = re.search(r"\s", window)
    if match is None:
        return window  # No boundary available; keep the raw slice.

    snapped = window[match.end():].lstrip()
    # If snapping consumed everything meaningful, fall back to the raw window.
    return snapped if snapped else window


def _recursive_split(
    text: str,
    separators: list[str],
    chunk_size: int,
    chunk_overlap: int,
    result: list[str],
) -> None:
    """
    Internal recursive function that does the actual splitting.
    
    Algorithm:
    1. Take the first (best) separator
    2. Split the text by that separator
    3. Merge pieces back together up to chunk_size
    4. If any piece is still too big, recurse with the next separator
    """
    if not text:
        return
    
    # Base case: text fits in a chunk
    if len(text) <= chunk_size:
        result.append(text)
        return
    
    # Try each separator (best first)
    separator = separators[0] if separators else ""
    remaining_separators = separators[1:] if len(separators) > 1 else [""]
    
    if separator:
        pieces = text.split(separator)
    else:
        # Last resort: split by character (for very long words)
        pieces = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size - chunk_overlap)]
        result.extend(pieces)
        return
    
    # Merge pieces back together, respecting chunk_size
    current_chunk = ""
    
    for piece in pieces:
        # If adding this piece would exceed chunk_size
        test_chunk = current_chunk + separator + piece if current_chunk else piece
        
        if len(test_chunk) <= chunk_size:
            current_chunk = test_chunk
        else:
            # Save the current chunk if it has content
            if current_chunk:
                result.append(current_chunk)
                
                # Start new chunk with overlap from previous chunk
                if chunk_overlap > 0 and len(current_chunk) > chunk_overlap:
                    overlap_text = _tail_overlap(current_chunk, chunk_overlap)
                    current_chunk = overlap_text + separator + piece
                else:
                    current_chunk = piece
            else:
                current_chunk = piece
            
            # If the current chunk is STILL too big, recurse with next separator
            if len(current_chunk) > chunk_size:
                _recursive_split(current_chunk, remaining_separators, chunk_size, chunk_overlap, result)
                current_chunk = ""
    
    # Don't forget the last chunk
    if current_chunk:
        if len(current_chunk) > chunk_size:
            _recursive_split(current_chunk, remaining_separators, chunk_size, chunk_overlap, result)
        else:
            result.append(current_chunk)


# ══════════════════════════════════════════════════════════════
# STAGE 3+4 — THE FULL INGESTION PIPELINE
# ══════════════════════════════════════════════════════════════
# Combines parsing, chunking, embedding, and storage into one
# function that takes a file in and puts knowledge out.
# ══════════════════════════════════════════════════════════════


async def ingest_document(
    file_bytes: bytes,
    filename: str,
    file_type: str,
    document_id: str,
    source_type: str = "general",
    title: str | None = None,
) -> dict:
    """
    THE MAIN PIPELINE — Process a document from raw bytes to stored vectors.
    
    This is called by the upload API endpoint. It runs the full 4-stage
    pipeline and returns a summary of what was done.
    
    Args:
        file_bytes: Raw bytes of the uploaded file
        filename: Original filename (for metadata)
        file_type: File extension without dot ("pdf", "png", etc.)
        document_id: UUID from PostgreSQL (links vectors back to metadata)
        source_type: Category ("textbook", "guideline", "statpearls", etc.)
        title: Optional document title
        
    Returns:
        Dict with processing results:
        {
            "status": "completed",
            "chunk_count": 45,
            "char_count": 23456,
            "message": "Successfully processed 45 chunks"
        }
    """
    # ── STAGE 1: Parse ──────────────────────────────────────
    # Parsing is CPU-bound (PyMuPDF, Tesseract) so it runs in a worker
    # thread, not on the event loop. See _embed_and_store for why.
    try:
        print(f"📄 Parsing {filename} ({file_type})...")
        raw_text = await run_in_threadpool(parse_file, file_bytes, file_type)
    except ParseError as e:
        # Expected, user-facing failure — report it cleanly.
        print(f"❌ Parse failed for {filename}: {e}")
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": 0,
            "message": str(e),
        }
    except Exception as e:
        # Unexpected — this is a bug, so log the full traceback.
        print(f"❌ Unexpected parse error for {filename}:")
        traceback.print_exc()
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": 0,
            "message": f"Internal error while parsing {filename}: {e}",
        }

    if len(raw_text.strip()) < MIN_USABLE_TEXT_CHARS:
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": len(raw_text),
            "message": (
                f"Only {len(raw_text.strip())} characters could be extracted, "
                f"which is below the {MIN_USABLE_TEXT_CHARS}-character minimum "
                f"for a useful knowledge base entry."
            ),
        }

    print(f"   → Extracted {len(raw_text)} characters")

    metadata = {"filename": filename, "source_type": source_type}
    if title:
        metadata["title"] = title

    return await _chunk_embed_store(
        text=raw_text,
        document_id=document_id,
        metadata=metadata,
        label=filename,
    )


async def ingest_text_content(
    text: str,
    title: str,
    document_id: str,
    source_type: str = "general",
    source_url: str | None = None,
) -> dict:
    """
    Ingest plain text content directly (not from a file).
    
    WHEN IS THIS USED?
    - Ingesting StatPearls articles fetched from NCBI API
    - Ingesting curated medical knowledge (seed data)
    - Ingesting text pasted by a doctor (no file upload needed)
    
    Skips Stage 1 (parsing) since we already have raw text.
    
    Args:
        text: The raw text content to ingest
        title: Title for the content
        document_id: UUID from PostgreSQL
        source_type: Category ("statpearls", "curated", etc.)
        source_url: Original URL if fetched from the internet
        
    Returns:
        Dict with processing results (same format as ingest_document)
    """
    if not text or len(text.strip()) < MIN_USABLE_TEXT_CHARS:
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": len(text or ""),
            "message": "Text content is too short or empty.",
        }

    metadata = {
        "filename": title,
        "source_type": source_type,
        "title": title,
    }
    if source_url:
        metadata["source_url"] = source_url

    return await _chunk_embed_store(
        text=text,
        document_id=document_id,
        metadata=metadata,
        label=title,
    )


# ══════════════════════════════════════════════════════════════
# SHARED BACK-HALF OF THE PIPELINE (Stages 2-4)
# ══════════════════════════════════════════════════════════════


async def _chunk_embed_store(
    text: str,
    document_id: str,
    metadata: dict,
    label: str,
) -> dict:
    """
    Chunk → embed → store. Shared by both ingest entry points.

    ⚠️  WHY run_in_threadpool?
    `embedding_service.encode()` is synchronous, CPU-bound PyTorch. Calling it
    directly inside an `async def` blocks the event loop — meaning the ENTIRE
    server stops serving requests (including /health) for as long as the
    embedding takes. On a 50 MB PDF that's minutes of total downtime.

    `run_in_threadpool` hands the work to a worker thread so the event loop
    stays free. PyTorch releases the GIL during its heavy math, so this
    genuinely runs in parallel rather than just moving the problem.

    We do the whole chunk+embed+store block in ONE thread hop rather than
    three, to avoid paying the context-switch cost repeatedly.
    """

    # ⚠️  Fail fast with a diagnosis, not a generic error.
    # If the embedding model didn't load at startup, EVERY document will fail.
    # Previously each one recorded "Internal error during ingestion: ..." with
    # the raw HuggingFace exception, which gave no hint that the real problem
    # was a missing model cache. Naming the cause once, clearly, turns a
    # baffling "34 documents failed" into an actionable message.
    if not embedding_service.is_loaded:
        msg = (
            "The embedding model is not loaded, so nothing can be indexed. "
            "This usually means the model cache was deleted (e.g. by "
            "`docker-compose down -v`) while HF_HUB_OFFLINE=1 blocked "
            "re-downloading it. Fix: HF_HUB_OFFLINE=0 docker-compose up -d backend"
        )
        print(f"❌ {label}: {msg}")
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": len(text),
            "message": msg,
        }

    def _contextualise(chunk: str) -> str:
        """
        Prefix a chunk with its document title before embedding.

        ⚠️  MEASURED NEED — this is the fix for the last stubborn failure.
        The chunk answering "What are the radiographic findings of
        pneumothorax?" reads:

            "CXR Findings of Pneumothorax: Visceral pleural line: A thin
             white line separated from the chest wall..."

        Neither retriever finds it. The embedding is dominated by "visceral
        pleural line" — vocabulary that appears in the ANSWER but not the
        QUESTION. BM25 fails for the same reason: the query terms
        (radiographic, findings, pneumothorax) are common across 248 papers
        that mention pneumothorax far more often.

        Prefixing the title gives every chunk its document-level context:

            "Pneumothorax — Types, Imaging, and Management. CXR Findings of
             Pneumothorax: Visceral pleural line: ..."

        Now the chunk itself contains "Pneumothorax", "Imaging" and
        "Findings", so it matches the question on BOTH retrievers. A chunk
        stripped of its source loses the very context that makes it findable.

        The prefix is used for EMBEDDING AND BM25 ONLY. The evidence panel
        still displays the original text — the title is already shown beside
        it, and repeating it inside the excerpt would just be noise.
        """
        label = (metadata.get("title") or metadata.get("filename") or "").strip()
        return f"{label}. {chunk}" if label else chunk

    def _blocking_work() -> tuple[list[str], int]:
        # ── STAGE 2: Chunk ──
        print(f"✂️  Chunking (size={settings.CHUNK_SIZE}, overlap={settings.CHUNK_OVERLAP})...")
        chunks = chunk_text(text)
        if not chunks:
            raise ParseError("Text was extracted but no valid chunks were produced.")
        print(f"   → Created {len(chunks)} chunks")

        # ── STAGE 3: Embed ──
        print(f"🧠 Embedding {len(chunks)} chunks (title-contextualised)...")
        search_texts = [_contextualise(c) for c in chunks]
        vectors = embedding_service.encode(
            search_texts,
            batch_size=settings.EMBEDDING_BATCH_SIZE,
            show_progress=len(chunks) > 50,
        )
        print(f"   → Generated {len(vectors)} vectors")

        # ── STAGE 4: Store ──
        # Delete first so re-ingesting a document replaces rather than
        # duplicates its chunks.
        print(f"💾 Storing in Qdrant collection '{settings.QDRANT_COLLECTION}'...")
        qdrant_service.delete_by_document(document_id)
        upserted = qdrant_service.upsert_chunks(
            document_id=document_id,
            chunks=chunks,
            vectors=vectors,
            metadata=metadata,
            search_texts=search_texts,
        )
        # New chunks exist now, so the BM25 snapshot is out of date.
        # The next lexical search rebuilds it.
        lexical_index.mark_stale()

        return chunks, upserted

    try:
        chunks, upserted = await run_in_threadpool(_blocking_work)
    except ParseError as e:
        print(f"❌ {label}: {e}")
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": len(text),
            "message": str(e),
        }
    except Exception as e:
        print(f"❌ Unexpected ingestion error for {label}:")
        traceback.print_exc()
        return {
            "status": "failed",
            "chunk_count": 0,
            "char_count": len(text),
            "message": f"Internal error during ingestion: {e}",
        }

    print(f"   → Stored {upserted} vectors in Qdrant")
    print(f"✅ Successfully ingested: {label}")

    return {
        "status": "completed",
        "chunk_count": len(chunks),
        "char_count": len(text),
        "message": f"Successfully processed {len(chunks)} chunks from '{label}'",
    }


# ══════════════════════════════════════════════════════════════
# RE-INDEXING
# ══════════════════════════════════════════════════════════════


async def reindex_all_chunks(batch_size: int = 256) -> dict:
    """
    Re-embed every stored chunk with the current embedding strategy.

    WHY IN-PLACE RATHER THAN RE-INGEST:
    Re-ingesting means re-fetching 200+ articles from NCBI — the single most
    failure-prone operation in this project on an unstable connection. The
    text is already in Qdrant; only the vectors need to change. This reads,
    re-embeds, and writes back under the same point IDs.

    Applies the title-prefix contextualisation (see _contextualise) so
    existing chunks gain the document context that new ingestions get
    automatically.
    """
    def _work() -> dict:
        records = qdrant_service.iter_all_chunks()
        if not records:
            return {"chunks": 0, "updated": 0, "message": "nothing to reindex"}

        total = len(records)
        print(f"🔁 Re-indexing {total:,} chunks with title context...")

        updated = 0
        for start in range(0, total, batch_size):
            batch = records[start : start + batch_size]

            # Prefix each chunk with its document title.
            search_texts = []
            for r in batch:
                label = (r.get("title") or r.get("filename") or "").strip()
                search_texts.append(f"{label}. {r['text']}" if label else r["text"])

            vectors = embedding_service.encode(
                search_texts, batch_size=settings.EMBEDDING_BATCH_SIZE
            )

            updated += qdrant_service.update_vectors(
                point_ids=[r["point_id"] for r in batch],
                vectors=vectors,
                search_texts=search_texts,
            )

            done = min(start + batch_size, total)
            print(f"   {done:,}/{total:,} ({100 * done // total}%)")

        lexical_index.mark_stale()
        print(f"✅ Re-index complete: {updated:,} chunks re-embedded")
        return {"chunks": total, "updated": updated}

    return await run_in_threadpool(_work)
